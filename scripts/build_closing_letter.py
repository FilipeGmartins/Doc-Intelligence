from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "CLOSING_LETTER.md"
OUTPUT = ROOT / "docs" / "CARTA_ENCERRAMENTO_DOC_INTELLIGENCE.docx"

BLUE = RGBColor(40, 91, 212)
DARK = RGBColor(28, 39, 57)
MUTED = RGBColor(105, 117, 136)
LIGHT_BLUE = "EDF3FF"


def set_font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Roboto"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Roboto")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Roboto")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Roboto"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Roboto")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Roboto")
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading = document.styles["Heading 1"]
    heading.font.name = "Roboto"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Roboto")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Roboto")
    heading.font.size = Pt(13)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.keep_with_next = True


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.start_type = WD_SECTION.NEW_PAGE
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("DOC Intelligence  |  Trilha B")
    set_font(header_run, size=8.5, bold=True, color=MUTED)
    add_page_field(section.footer.paragraphs[0])

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(5)
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("CARTA DE ENCERRAMENTO")
    set_font(title_run, size=20, bold=True, color=DARK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("DOC Intelligence - Processamento e conferência assistida de documentos")
    set_font(subtitle_run, size=11.5, bold=True, color=BLUE)

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(12)
    metadata_run = metadata.add_run("Escopo: Front-end com API, banco e IA simulados  |  Setembro de 2026")
    set_font(metadata_run, size=9, color=MUTED)

    lead = document.add_paragraph()
    lead.paragraph_format.left_indent = Inches(0.16)
    lead.paragraph_format.right_indent = Inches(0.16)
    lead.paragraph_format.space_before = Pt(0)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.15
    shade_paragraph(lead, LIGHT_BLUE)
    lead_run = lead.add_run(
        "A entrega prioriza uma fatia vertical pequena, funcional e substituível: "
        "recebimento por cliente, processamento determinístico, conferência humana, "
        "auditoria e consulta, sem uso de dados ou integrações reais."
    )
    set_font(lead_run, size=10.5, bold=True, color=RGBColor(46, 73, 122))

    blocks = [block.strip() for block in SOURCE.read_text(encoding="utf-8").split("\n\n") if block.strip()]
    skip_overview = True
    for block in blocks:
        if block.startswith("# Carta"):
            continue
        if block == "## Visão geral":
            skip_overview = True
            continue
        if block.startswith("## "):
            skip_overview = False
            document.add_paragraph(block[3:], style="Heading 1")
            continue
        if skip_overview:
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        text = " ".join(line.strip() for line in block.splitlines())
        run = paragraph.add_run(text.replace("`", ""))
        set_font(run)

    document.core_properties.title = "Carta de encerramento - DOC Intelligence"
    document.core_properties.subject = "Desafio técnico - Trilha B"
    document.core_properties.author = "Filipe Gomes Martins"
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
